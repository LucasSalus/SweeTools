import os
import json
import qrcode
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkinter import filedialog, messagebox

CONFIG_FILE = "gateway_config.json"

def obter_memorias():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r") as f:
                return json.load(f)
        except:
            return {}
    return {}

def salvar_memoria(chave, caminho):
    memorias = obter_memorias()
    memorias[chave] = caminho
    try:
        with open(CONFIG_FILE, "w") as f:
            json.dump(memorias, f)
    except:
        pass

def escolher_nova_pasta(chave):
    """Esta é a função que estava faltando no seu erro!"""
    memorias = obter_memorias()
    diretorio_inicial = memorias.get(chave)
    pasta = filedialog.askdirectory(title="Selecionar Pasta de Destino", initialdir=diretorio_inicial)
    if pasta:
        salvar_memoria(chave, pasta)
        return pasta
    return None

def processar_qr(url, nome, pasta_destino):
    if not url or not nome:
        messagebox.showwarning("Erro", "Preencha a URL e o Nome!")
        return False
    if not pasta_destino or pasta_destino == "Não definido":
        messagebox.showwarning("Atenção", "Selecione a pasta de destino antes!")
        return False

    try:
        caminho_final = os.path.join(pasta_destino, f"{nome}.png")
        qrcode.make(url).save(caminho_final)
        return messagebox.askyesno("Sucesso", f"QR '{nome}' gerado!\nDeseja realizar outro?")
    except Exception as e:
        messagebox.showerror("Erro", str(e))
        return False

def processar_csv_para_word(nome_base, pasta_destino):
    if not nome_base or not pasta_destino or pasta_destino == "Não definido":
        messagebox.showwarning("Erro", "Verifique o nome e o destino antes!")
        return False
    
    caminho_csv = filedialog.askopenfilename(title="Selecione o CSV", filetypes=[("CSV", "*.csv")])
    if not caminho_csv: return False

    try:
        try: df = pd.read_csv(caminho_csv, sep=';')
        except: df = pd.read_csv(caminho_csv, sep=';', encoding='latin-1')
        
        df = df.fillna('-')
        contador = 1
        for i in range(0, len(df), 1000):
            bloco = df.iloc[i : i + 1000]
            doc = Document()
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

            doc.add_heading(f'{nome_base.replace("_", " ")} - Parte {contador}', 0)
            tabela = doc.add_table(rows=1, cols=len(bloco.columns))
            tabela.style = 'Table Grid'

            for j, col in enumerate(bloco.columns):
                c = tabela.rows[0].cells[j]
                c.text = str(col)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            for _, linha in bloco.iterrows():
                celulas = tabela.add_row().cells
                for k, v in enumerate(linha):
                    v_str = str(v)
                    celulas[k].text = v_str
                    if v_str == '-': celulas[k].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(os.path.join(pasta_destino, f'{nome_base}_{contador}.docx'))
            contador += 1
        
        return messagebox.askyesno("Sucesso", f"Concluído! {contador-1} arquivos gerados.\nDeseja realizar outro?")
    except Exception as e:
        messagebox.showerror("Erro", str(e))
        return False